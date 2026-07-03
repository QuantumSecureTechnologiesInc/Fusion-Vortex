"""
Extract Rust source from .docx files in "Fusion Crates" folder,
apply Rust→Fusion type conversions, and place into registry/crates/.
"""
import docx
import os
import re
import shutil
from pathlib import Path

DOCX_DIR = r"C:\Users\Matth\Downloads\New folder\Fusion Crates"
WORKSPACE = r"C:\Users\Matth\Downloads\Fusion v2.0 Vortex"

# ── Crate routing ──────────────────────────────────────────────────
# Files go to registry/crates/<name>/src/  (default) or crates/<name>/src/ (override)
REGISTRY_OVERRIDE = {
    "ai-core": "crates/ai-core",  # existing crate under crates/, not registry
}

# Which crates are NEW (don't exist yet in registry or crates)? None in registry yet.
NEW_CRATES = {
    "ai-agents", "ai-gan", "ai-hf-transformers",
    "llm-builder", "llm-chain", "llm-eval",
    "rl-gym", "toolchain",
}

KNOWN_REGISTRY_CRATES = set()  # populated at runtime by scanning

def scan_registry():
    """Find all existing crate names in registry/crates/ and crates/."""
    existing = set()
    for base in ["registry/crates", "crates"]:
        p = Path(WORKSPACE) / base
        if p.exists():
            for d in p.iterdir():
                if d.is_dir() and not d.name.startswith("."):
                    existing.add(d.name)
    return existing

def rust_to_fusion_transform(text):
    """Apply Rust→Fusion type/name substitutions."""
    # Ordered replacements — longest first to avoid partial matches
    replacements = [
        # Type aliases
        (r'\bVec<', 'FVec<'),
        (r'\bString\b', 'FString'),
        (r'\busize\b', 'FSize'),
        (r'\bisize\b', 'FInt'),
        (r'\bu8\b', 'u8'),         # keep
        (r'\bu16\b', 'u16'),
        (r'\bu32\b', 'u32'),
        (r'\bu64\b', 'u64'),
        (r'\bi8\b', 'i8'),
        (r'\bi16\b', 'i16'),
        (r'\bi32\b', 'i32'),
        (r'\bi64\b', 'i64'),
        (r'\bf32\b', 'f32'),
        (r'\bf64\b', 'f64'),
        (r'\bbool\b', 'bool'),
        # Box/Arc/Rc
        (r'\bBox<', 'FBox<'),
        (r'\bArc<', 'FArc<'),
        (r'\bRc<', 'FRc<'),
        # HashMap/HashSet/BTreeMap
        (r'\bHashMap<', 'FHashMap<'),
        (r'\bHashSet<', 'FHashSet<'),
        (r'\bBTreeMap<', 'FBTreeMap<'),
        (r'\bBTreeSet<', 'FBTreeSet<'),
    ]
    for pattern, replacement in replacements:
        text = re.sub(pattern, replacement, text)
    return text

def parse_docx_filename(filename):
    """
    Parse 'crates_ai-agents_src_runtime.rs.docx' → ('ai-agents', 'src', 'runtime.rs')
    Parse 'crates_ai-core_src_layers_conv.rs.docx' → ('ai-core', 'src/layers', 'conv.rs')
    """
    name = filename.replace('.docx', '')
    # Remove 'crates_' prefix
    if name.startswith('crates_'):
        name = name[7:]
    # Split into crate name and remaining path
    # Pattern: <crate>_src_<rest>.rs  or  <crate>_src_<dir>_<file>.rs
    parts = name.split('_src_', 1)
    if len(parts) != 2:
        return None
    crate_name = parts[0].replace('_', '-')  # ai_core → ai-core
    rest = parts[1]  # e.g., 'layers_conv.rs' or 'lib.rs' or 'runtime.rs'
    
    # Split rest into dir parts + filename
    rest_parts = rest.split('_')
    if len(rest_parts) == 1:
        filename = rest_parts[0]  # e.g., 'lib.rs'
        subdir = ''
    else:
        filename = rest_parts[-1]  # last part is filename
        subdir_parts = rest_parts[:-1]
        subdir = '/'.join(subdir_parts)  # e.g., 'layers'
    
    return (crate_name, subdir, filename)

def get_dest_path(crate_name, subdir, filename, existing_creates):
    """Determine the destination directory for a crate."""
    base = REGISTRY_OVERRIDE.get(crate_name)
    if base:
        dest = Path(WORKSPACE) / base / "src"
    elif crate_name in existing_creates:
        dest = Path(WORKSPACE) / "registry" / "crates" / crate_name / "src"
    else:
        # New crate
        dest = Path(WORKSPACE) / "registry" / "crates" / crate_name / "src"
    
    if subdir:
        dest = dest / subdir
    
    return dest

def create_fusion_toml(crate_name: str, dest_dir: Path):
    """Create a basic Fusion.toml for a new crate."""
    toml_path = dest_dir.parent / "Fusion.toml"
    if toml_path.exists():
        return
    
    description_map = {
        "ai-agents": "Agent: AI agent runtime with LLM integration and tool orchestration",
        "ai-gan": "Algorithm: Generative Adversarial Network components",
        "ai-hf-transformers": "Integration: HuggingFace Transformers model loader and bridge",
        "llm-builder": "Algorithm: LLM computation graph builder and optimizer",
        "llm-chain": "Algorithm: LLM chain executor for multi-step reasoning",
        "llm-eval": "Tool: LLM evaluation runner and benchmarking framework",
        "rl-gym": "Algorithm: Reinforcement Learning environment and algorithms",
        "toolchain": "Tool: Fusion build toolchain and compilation pipeline",
    }
    
    desc = description_map.get(crate_name, f"Fusion crate: {crate_name}")
    keywords = crate_name.replace('-', ' ')
    
    toml = f"""[package]
name = "{crate_name}"
version = "0.1.0"
edition = "2021"
license = "MIT"
description = "{desc}"
authors = ["Fusion Team"]
repository = "https://github.com/QuantumSecureTechnologiesInc/Fusion-Programming-Language"
keywords = ["{keywords}"]
categories = ["algorithms"]

[dependencies]
fusion_core = {{ workspace = true }}
fusion_std = "1.0.0"
"""
    toml_path.parent.mkdir(parents=True, exist_ok=True)
    toml_path.write_text(toml, encoding='utf-8')

def main():
    existing = scan_registry()
    print(f"Existing crates: {len(existing)}")
    
    docx_files = list(Path(DOCX_DIR).glob("*.docx"))
    print(f"Found {len(docx_files)} .docx files")
    
    created_crates = set()
    written_files = []
    skipped = []
    errors = []
    
    for docx_path in docx_files:
        fname = docx_path.name
        
        # Skip workspace config files
        if "Updated Workspace" in fname:
            print(f"  SKIP (workspace): {fname}")
            skipped.append(fname)
            continue
        
        # Parse filename
        parsed = parse_docx_filename(fname)
        if not parsed:
            print(f"  SKIP (parse fail): {fname}")
            skipped.append(fname)
            continue
        
        crate_name, subdir, filename = parsed
        
        # Extract text from .docx
        try:
            doc = docx.Document(str(docx_path))
            if len(doc.paragraphs) == 0:
                errors.append(f"Empty docx: {fname}")
                continue
            text = doc.paragraphs[0].text
        except Exception as e:
            errors.append(f"Read error {fname}: {e}")
            continue
        
        if not text.strip():
            errors.append(f"Empty text: {fname}")
            continue
        
        # Determine destination
        dest_dir = get_dest_path(crate_name, subdir, filename, existing)
        
        # Create directories
        dest_dir.mkdir(parents=True, exist_ok=True)
        
        # Determine output filename
        base_name = filename.replace('.rs', '')
        rs_output = dest_dir / f"{base_name}.rs"
        fu_output = dest_dir / f"{base_name}.fu"
        
        # Write original Rust
        rs_output.write_text(text, encoding='utf-8')
        
        # Transform and write .fu
        fu_text = rust_to_fusion_transform(text)
        fu_output.write_text(fu_text, encoding='utf-8')
        
        # Create Fusion.toml for new crates
        if crate_name in NEW_CRATES and crate_name not in existing:
            create_fusion_toml(crate_name, dest_dir)
            created_crates.add(crate_name)
        
        display = f"{crate_name}/{subdir}/{base_name}" if subdir else f"{crate_name}/{base_name}"
        print(f"  OK: {display}.fu ({len(fu_text)} chars)")
        written_files.append(str(fu_output.relative_to(WORKSPACE)))
    
    # Summary
    print(f"\n{'='*60}")
    print(f"Written: {len(written_files)} .fu files")
    print(f"New crates: {len(created_crates)}")
    print(f"Skipped: {len(skipped)}")
    print(f"Errors: {len(errors)}")
    
    if created_crates:
        print(f"\nNew crates created:")
        for c in sorted(created_crates):
            print(f"  - registry/crates/{c}/")
    
    if errors:
        print(f"\nErrors:")
        for e in errors:
            print(f"  {e}")

if __name__ == "__main__":
    main()